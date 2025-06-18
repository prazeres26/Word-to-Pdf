from flask import Flask, request, render_template, send_file
import os
import io
from werkzeug.utils import secure_filename
import fitz  # PyMuPDF
from docx import Document
from fpdf import FPDF

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = '/tmp'

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

def docx_to_pdf(docx_file):
    doc = Document(docx_file)
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    for para in doc.paragraphs:
        pdf.multi_cell(0, 10, para.text)
    output = io.BytesIO()
    pdf.output(output)
    output.seek(0)
    return output

def pdf_to_text(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    output = io.StringIO()
    for page in doc:
        output.write(page.get_text())
    doc.close()
    text = output.getvalue()
    docx_io = io.BytesIO()
    docx_doc = Document()
    docx_doc.add_paragraph(text)
    docx_doc.save(docx_io)
    docx_io.seek(0)
    return docx_io

@app.route('/convert', methods=['POST'])
def convert():
    file = request.files.get('file')
    convert_to = request.form.get('convert_to', 'pdf')
    if not file:
        return 'Arquivo inválido', 400
    filename = secure_filename(file.filename)
    try:
        if convert_to == 'pdf' and filename.lower().endswith('.docx'):
            result = docx_to_pdf(file)
            return send_file(result, mimetype='application/pdf', download_name='convertido.pdf')
        elif convert_to == 'docx' and filename.lower().endswith('.pdf'):
            result = pdf_to_text(file)
            return send_file(result, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                             download_name='convertido.docx')
        else:
            return 'Conversão não suportada para este tipo de arquivo.', 400
    except Exception as e:
        return f'Erro ao converter: {str(e)}', 500